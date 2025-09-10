from flask import Flask, render_template, request, jsonify, send_file, session
from transformers import pipeline
import PyPDF2
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import uuid
import threading

app = Flask(__name__)
app.secret_key = "supersecretkey"  # Required for session

# Load summarization pipeline
summarizer = pipeline("summarization")

progress_store = {}
progress_lock = threading.Lock()

def chunk_text(text, max_tokens=2000):
    words = text.split()
    chunks = []
    current_chunk = []
    count = 0
    for word in words:
        current_chunk.append(word)
        count += 1
        if count >= max_tokens:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            count = 0
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

# Text extraction functions
def extract_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + " "
    return text

def extract_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + " "
    return text

def extract_txt(file):
    return file.read().decode("utf-8")

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/start_summarization", methods=["POST"])
def start_summarization():
    uploaded_file = request.files.get("file")
    text_input = request.form.get("text")

    text_to_summarize = ""
    if uploaded_file and uploaded_file.filename != "":
        filename = uploaded_file.filename.lower()
        if filename.endswith(".pdf"):
            text_to_summarize = extract_pdf(uploaded_file)
        elif filename.endswith(".docx"):
            text_to_summarize = extract_docx(uploaded_file)
        elif filename.endswith(".txt"):
            text_to_summarize = extract_txt(uploaded_file)
    elif text_input:
        text_to_summarize = text_input

    if not text_to_summarize.strip():
        return jsonify({"error": "No text provided."}), 400

    # Prepare unique session ID for progress tracking
    session_id = str(uuid.uuid4())
    session['session_id'] = session_id

    # Initialize progress
    with progress_lock:
        progress_store[session_id] = {"progress": 0, "summary": ""}

    # Start summarization asynchronously
    from threading import Thread

    def summarize_chunks():
        import time
        chunks = chunk_text(text_to_summarize, max_tokens=50)
        total_chunks = len(chunks)
        print(f"Number of chunks: {total_chunks}")
        summaries = []
        for i, chunk in enumerate(chunks):
            summary_chunk = summarizer(chunk, max_length=15, min_length=5, do_sample=False)[0]['summary_text']
            summaries.append(summary_chunk)
            with progress_lock:
                progress_store[session_id]['progress'] = int(((i + 1)/total_chunks) * 100)
                progress_store[session_id]['summary'] = " ".join(summaries)
            print(f"Progress for {session_id}: {progress_store[session_id]['progress']}%")
            # Removed artificial delay for faster processing

    thread = Thread(target=summarize_chunks)
    thread.start()

    return jsonify({"session_id": session_id})

@app.route("/progress/<session_id>")
def get_progress(session_id):
    from flask import make_response
    with progress_lock:
        data = progress_store.get(session_id, {"progress": 0, "summary": ""})
    response = make_response(jsonify(data))
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

@app.route("/download", methods=["POST"])
def download():
    summary = request.form.get("summary")
    filetype = request.form.get("filetype")
    if not summary:
        return "No summary to download.", 400

    if filetype == "txt":
        return send_file(io.BytesIO(summary.encode()), download_name="summary.txt", as_attachment=True)
    elif filetype == "pdf":
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        textobject = c.beginText(40, 750)
        lines = summary.split('\n')
        for line in lines:
            textobject.textLine(line)
        c.drawText(textobject)
        c.save()
        buffer.seek(0)
        return send_file(buffer, download_name="summary.pdf", as_attachment=True)
    elif filetype == "docx":
        doc = docx.Document()
        doc.add_paragraph(summary)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, download_name="summary.docx", as_attachment=True)
    else:
        return "Unsupported file type.", 400

if __name__ == "__main__":
    app.run(debug=True, threaded=True)
